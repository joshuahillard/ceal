"""
Ceal Phase 2: Resume Export to .docx

Exports tailored resume bullets to a formatted Word document.
Each job gets a section with the company name, title, and rewritten bullets.
"""
from __future__ import annotations

from pathlib import Path

import structlog
from docx import Document
from docx.shared import Pt

from src.tailoring.models import TailoringResult

logger = structlog.get_logger(__name__)


def export_tailoring_result(
    result: TailoringResult,
    job_title: str,
    company_name: str,
    output_path: str | Path,
) -> Path:
    """
    Export a single TailoringResult to a .docx file.

    Args:
        result: The tailoring result with rewritten bullets and skill gaps.
        job_title: Job title for the document header.
        company_name: Company name for the document header.
        output_path: Where to save the .docx file.

    Returns:
        Path to the saved .docx file.
    """
    output_path = Path(output_path)
    doc = Document()

    # Title
    doc.add_heading(f"Tailored Resume — {company_name}", level=1)
    doc.add_paragraph(f"Position: {job_title}")
    doc.add_paragraph(f"Tier: {result.request.target_tier}")

    if result.request.emphasis_areas:
        doc.add_paragraph(
            f"Emphasis: {', '.join(result.request.emphasis_areas)}"
        )

    # Tailored bullets
    doc.add_heading("Tailored Bullets", level=2)
    for bullet in result.tailored_bullets:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(bullet.rewritten_text)
        run.font.size = Pt(11)

        if bullet.xyz_format:
            tag = doc.add_paragraph()
            tag_run = tag.add_run("[X-Y-Z format]")
            tag_run.italic = True
            tag_run.font.size = Pt(9)

    # Skill gaps summary
    if result.skill_gaps:
        doc.add_heading("Skill Gap Analysis", level=2)
        matched = [g for g in result.skill_gaps if g.resume_has]
        missing = [g for g in result.skill_gaps if not g.resume_has]

        if matched:
            doc.add_paragraph("Matched Skills:")
            for gap in matched:
                doc.add_paragraph(
                    f"{gap.skill_name} ({gap.category.value})",
                    style="List Bullet",
                )

        if missing:
            doc.add_paragraph("Missing Skills:")
            for gap in missing:
                doc.add_paragraph(
                    f"{gap.skill_name} ({gap.category.value})",
                    style="List Bullet",
                )

    doc.save(str(output_path))
    logger.info(
        "export_saved",
        path=str(output_path),
        bullets=len(result.tailored_bullets),
    )
    return output_path
