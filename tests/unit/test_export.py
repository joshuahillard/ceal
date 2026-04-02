"""
Ceal Phase 2: Export Tests

Tests .docx generation from TailoringResult data.
"""
from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document

from src.export import export_tailoring_result
from src.models.entities import Proficiency, SkillCategory
from src.tailoring.models import (
    SkillGap,
    TailoredBullet,
    TailoringRequest,
    TailoringResult,
)


def _make_result(**overrides) -> TailoringResult:
    """Build a minimal TailoringResult for export tests."""
    defaults = {
        "request": TailoringRequest(
            job_id=1, profile_id=1, target_tier=1, emphasis_areas=["Python"],
        ),
        "tailored_bullets": [
            TailoredBullet(
                original="Built REST APIs",
                rewritten_text=(
                    "Accomplished 40% latency reduction as measured by "
                    "p99 response times, by doing async refactor of REST endpoints"
                ),
                xyz_format=True,
                relevance_score=0.9,
            ),
        ],
        "skill_gaps": [
            SkillGap(
                skill_name="Python",
                category=SkillCategory.LANGUAGE,
                job_requires=True,
                resume_has=True,
                proficiency=Proficiency.EXPERT,
            ),
            SkillGap(
                skill_name="Kubernetes",
                category=SkillCategory.TOOL,
                job_requires=True,
                resume_has=False,
            ),
        ],
        "tailoring_version": "v1.0",
    }
    defaults.update(overrides)
    return TailoringResult(**defaults)


class TestExportDocx:
    """Tests for .docx export."""

    def test_export_creates_file(self):
        """Export produces a .docx file on disk."""
        result = _make_result()
        with tempfile.TemporaryDirectory() as tmpdir:
            out = export_tailoring_result(
                result=result,
                job_title="TSE",
                company_name="Acme",
                output_path=Path(tmpdir) / "test.docx",
            )
            assert out.exists()
            assert out.suffix == ".docx"

    def test_export_contains_bullets(self):
        """Exported doc contains the rewritten bullet text."""
        result = _make_result()
        with tempfile.TemporaryDirectory() as tmpdir:
            out = export_tailoring_result(
                result=result,
                job_title="TSE",
                company_name="Acme",
                output_path=Path(tmpdir) / "test.docx",
            )
            doc = Document(str(out))
            full_text = "\n".join(p.text for p in doc.paragraphs)
            assert "latency reduction" in full_text

    def test_export_contains_skill_gaps(self):
        """Exported doc includes matched and missing skill sections."""
        result = _make_result()
        with tempfile.TemporaryDirectory() as tmpdir:
            out = export_tailoring_result(
                result=result,
                job_title="TSE",
                company_name="Acme",
                output_path=Path(tmpdir) / "test.docx",
            )
            doc = Document(str(out))
            full_text = "\n".join(p.text for p in doc.paragraphs)
            assert "Python" in full_text
            assert "Kubernetes" in full_text

    def test_export_contains_header_info(self):
        """Exported doc has company name, job title, and tier."""
        result = _make_result()
        with tempfile.TemporaryDirectory() as tmpdir:
            out = export_tailoring_result(
                result=result,
                job_title="Senior Engineer",
                company_name="BigCorp",
                output_path=Path(tmpdir) / "test.docx",
            )
            doc = Document(str(out))
            full_text = "\n".join(p.text for p in doc.paragraphs)
            assert "BigCorp" in full_text
            assert "Senior Engineer" in full_text

    def test_export_no_skill_gaps(self):
        """Export works when there are no skill gaps."""
        result = _make_result(skill_gaps=[])
        with tempfile.TemporaryDirectory() as tmpdir:
            out = export_tailoring_result(
                result=result,
                job_title="TSE",
                company_name="Acme",
                output_path=Path(tmpdir) / "test.docx",
            )
            doc = Document(str(out))
            full_text = "\n".join(p.text for p in doc.paragraphs)
            assert "Skill Gap" not in full_text
